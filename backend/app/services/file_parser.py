"""
File Parser Service
Extracts text from PDF and DOCX files
"""

from typing import Dict, Any, Optional
from loguru import logger
import io
import re

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not installed - PDF parsing disabled")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed - DOCX parsing disabled")


class FileParser:
    """Parse resume files (PDF, DOCX, TXT)"""

    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

    @staticmethod
    def extract_text_from_pdf(file_bytes: bytes) -> Dict[str, Any]:
        """
        Extract text from PDF file

        Args:
            file_bytes: PDF file content as bytes

        Returns:
            Dict with 'text' and 'metadata'
        """
        if not PDF_AVAILABLE:
            return {
                "error": "PDF parsing not available. Install PyPDF2: pip install PyPDF2",
                "text": ""
            }

        try:
            # Check file size
            if len(file_bytes) > FileParser.MAX_FILE_SIZE:
                return {
                    "error": "File too large (max 10MB)",
                    "text": ""
                }

            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            # Extract text from all pages
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num}: {e}")
                    continue

            full_text = "\n".join(text_parts)

            # Clean up text
            full_text = FileParser._clean_text(full_text)

            metadata = {
                "num_pages": len(pdf_reader.pages),
                "num_characters": len(full_text),
                "file_type": "pdf"
            }

            logger.info(f"✅ Extracted {len(full_text)} characters from PDF ({metadata['num_pages']} pages)")

            return {
                "text": full_text,
                "metadata": metadata
            }

        except Exception as e:
            logger.error(f"❌ PDF extraction failed: {e}")
            return {
                "error": f"Failed to parse PDF: {str(e)}",
                "text": ""
            }

    @staticmethod
    def extract_text_from_docx(file_bytes: bytes) -> Dict[str, Any]:
        """
        Extract text from DOCX file

        Args:
            file_bytes: DOCX file content as bytes

        Returns:
            Dict with 'text' and 'metadata'
        """
        if not DOCX_AVAILABLE:
            return {
                "error": "DOCX parsing not available. Install python-docx: pip install python-docx",
                "text": ""
            }

        try:
            # Check file size
            if len(file_bytes) > FileParser.MAX_FILE_SIZE:
                return {
                    "error": "File too large (max 10MB)",
                    "text": ""
                }

            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Extract tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_text.append(row_text)

            # Combine all text
            full_text = "\n".join(paragraphs)
            if table_text:
                full_text += "\n\n" + "\n".join(table_text)

            # Clean up text
            full_text = FileParser._clean_text(full_text)

            metadata = {
                "num_paragraphs": len(paragraphs),
                "num_tables": len(doc.tables),
                "num_characters": len(full_text),
                "file_type": "docx"
            }

            logger.info(f"✅ Extracted {len(full_text)} characters from DOCX")

            return {
                "text": full_text,
                "metadata": metadata
            }

        except Exception as e:
            logger.error(f"❌ DOCX extraction failed: {e}")
            return {
                "error": f"Failed to parse DOCX: {str(e)}",
                "text": ""
            }

    @staticmethod
    def extract_text_from_txt(file_bytes: bytes) -> Dict[str, Any]:
        """
        Extract text from TXT file

        Args:
            file_bytes: TXT file content as bytes

        Returns:
            Dict with 'text' and 'metadata'
        """
        try:
            # Check file size
            if len(file_bytes) > FileParser.MAX_FILE_SIZE:
                return {
                    "error": "File too large (max 10MB)",
                    "text": ""
                }

            # Try UTF-8 first, then fallback to latin-1
            try:
                text = file_bytes.decode('utf-8')
            except UnicodeDecodeError:
                text = file_bytes.decode('latin-1')

            # Clean up text
            text = FileParser._clean_text(text)

            metadata = {
                "num_characters": len(text),
                "file_type": "txt"
            }

            logger.info(f"✅ Extracted {len(text)} characters from TXT")

            return {
                "text": text,
                "metadata": metadata
            }

        except Exception as e:
            logger.error(f"❌ TXT extraction failed: {e}")
            return {
                "error": f"Failed to parse TXT: {str(e)}",
                "text": ""
            }

    @staticmethod
    def parse_file(file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse file based on extension

        Args:
            file_bytes: File content as bytes
            filename: Original filename

        Returns:
            Dict with 'text' and 'metadata'
        """
        # Detect file type from extension
        filename_lower = filename.lower()

        if filename_lower.endswith('.pdf'):
            return FileParser.extract_text_from_pdf(file_bytes)
        elif filename_lower.endswith('.docx'):
            return FileParser.extract_text_from_docx(file_bytes)
        elif filename_lower.endswith('.txt'):
            return FileParser.extract_text_from_txt(file_bytes)
        else:
            # Try to parse as text anyway
            logger.warning(f"Unknown file type: {filename}, attempting text extraction")
            return FileParser.extract_text_from_txt(file_bytes)

    @staticmethod
    def _clean_text(text: str) -> str:
        """
        Clean and normalize extracted text

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)

        # Remove common PDF artifacts
        text = re.sub(r'[^\S\n]+\n', '\n', text)

        # Normalize line endings
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # Strip leading/trailing whitespace from lines
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(lines)

        return text.strip()

    @staticmethod
    def validate_resume_content(text: str) -> Dict[str, Any]:
        """
        Validate that the extracted text looks like a resume

        Args:
            text: Extracted text

        Returns:
            Dict with validation results
        """
        validation = {
            "is_valid": True,
            "confidence": 1.0,
            "warnings": [],
            "has_contact_info": False,
            "has_experience": False,
            "has_education": False
        }

        # Check minimum length
        if len(text) < 100:
            validation["warnings"].append("Text is very short (< 100 characters)")
            validation["confidence"] = 0.3
            validation["is_valid"] = False
            return validation

        text_lower = text.lower()

        # Check for contact info
        if any(indicator in text_lower for indicator in ['email', '@', 'phone', 'linkedin']):
            validation["has_contact_info"] = True

        # Check for experience section
        experience_keywords = ['experience', 'work history', 'employment', 'professional']
        if any(keyword in text_lower for keyword in experience_keywords):
            validation["has_experience"] = True

        # Check for education section
        education_keywords = ['education', 'university', 'college', 'degree', 'bachelor', 'master']
        if any(keyword in text_lower for keyword in education_keywords):
            validation["has_education"] = True

        # Adjust confidence based on found sections
        sections_found = sum([
            validation["has_contact_info"],
            validation["has_experience"],
            validation["has_education"]
        ])

        if sections_found == 0:
            validation["is_valid"] = False
            validation["confidence"] = 0.2
            validation["warnings"].append("No typical resume sections found")
        elif sections_found == 1:
            validation["confidence"] = 0.5
            validation["warnings"].append("Only one resume section found")
        elif sections_found == 2:
            validation["confidence"] = 0.8
        else:
            validation["confidence"] = 1.0

        return validation


# Global instance
file_parser = FileParser()
